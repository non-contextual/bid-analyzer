"""
Word 文档解析器，支持 .docx（现代格式）和 .doc（Word 97-2003 旧格式）。

处理流程：
  1. 先用 python-docx 尝试解析（适用于 .docx）
  2. 失败时回退到 OLE 二进制文本提取（适用于 .doc）
"""

import io
import re

from parsers.pdf_parser import MAX_CHARS  # 复用同一截断长度


def parse_word(file_bytes: bytes) -> str:
    """
    从 Word 字节流中提取文本，自动判断 .docx / .doc 格式。

    Returns:
        提取的纯文本，截断到 MAX_CHARS 字符

    Raises:
        ValueError: 文件无法解析（损坏、格式不对等）
    """
    # ── 优先尝试 python-docx（处理 .docx 和部分兼容的 .doc）──────
    try:
        text = _parse_docx(file_bytes)
        if text.strip():
            return text
    except Exception:
        pass

    # ── 回退：OLE 二进制提取（.doc Word 97-2003）─────────────────
    try:
        text = _parse_doc_legacy(file_bytes)
        if text.strip():
            return text
    except Exception:
        pass

    raise ValueError(
        "无法读取 Word 文件，请检查文件是否损坏。"
        "如果是 .doc 格式，建议用 Word 另存为 .docx 后重新上传。"
    )


def _parse_docx(file_bytes: bytes) -> str:
    """用 python-docx 提取 .docx 文本（含表格）。"""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # 表格内容（资质要求、技术参数表格）
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("| " + " | ".join(cells) + " |")

    text = "\n".join(parts)
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[文档内容已截断，以上为前 8000 字]"
    return text.strip()


def _parse_doc_legacy(file_bytes: bytes) -> str:
    """
    对 .doc（OLE 二进制）格式做启发式文本提取。
    OLE2 文件以 magic bytes D0 CF 11 E0 开头，先验证格式再提取。
    """
    # OLE2 magic：Word 97-2003 .doc 的必要前缀
    OLE_MAGIC = b"\xd0\xcf\x11\xe0"
    if not file_bytes.startswith(OLE_MAGIC):
        raise ValueError("不是有效的 .doc 文件（缺少 OLE2 文件头）")

    # .doc 正文以 UTF-16 LE 存储，直接解码再提取可读片段
    raw = file_bytes.decode("utf-16-le", errors="ignore")

    # 只匹配中文字符及常见标点，过滤掉二进制噪声
    # 要求连续片段长度 >= 8（避免随机二进制碰巧通过）
    CJK = (
        "\u4e00-\u9fff"   # CJK 统一汉字
        "\u3000-\u303f"   # CJK 标点
        "\uff00-\uffef"   # 全角字符
    )
    chunks = re.findall(
        r"[" + CJK + r"\w\s]{8,}",
        raw,
    )

    text = "\n".join(c.strip() for c in chunks if c.strip())
    if not text:
        raise ValueError("无法从 .doc 文件中提取文本，建议另存为 .docx 后重新上传。")
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[文档内容已截断，以上为前 8000 字]"
    return text.strip()
